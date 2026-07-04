"""文档加载器 — 支持 txt/md/pdf/docx/csv/json，输出统一 Document 对象。

所有加载器均为 best-effort：解析失败时抛出 ValueError，调用方捕获后跳过该文件。
"""

from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# 支持的文件扩展名（小写，含点）
SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx", ".csv", ".json"}


@dataclass
class Document:
    """统一文档对象。"""

    id: str  # 文档 ID（通常为文件名 stem 或自定义）
    content: str  # 纯文本内容
    source: str  # 原始文件路径或 URL
    filename: str  # 文件名
    file_type: str  # 扩展名（不含点，如 "pdf"）
    size: int  # 文件大小（字节）
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    metadata: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "content": self.content,
            "source": self.source,
            "filename": self.filename,
            "file_type": self.file_type,
            "size": self.size,
            "created_at": self.created_at,
            "metadata": self.metadata,
        }


def load_document(file_path: str | Path, doc_id: Optional[str] = None) -> Document:
    """加载文档文件，返回统一 Document 对象。

    Args:
        file_path: 文件路径
        doc_id: 可选文档 ID，默认使用文件名 stem

    Returns:
        Document 对象

    Raises:
        ValueError: 不支持的文件类型或解析失败
    """
    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"文件不存在: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext}（支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}）"
        )

    file_type = ext.lstrip(".")
    doc_id = doc_id or path.stem
    size = path.stat().st_size

    if ext in (".txt",):
        content = _load_txt(path)
    elif ext in (".md", ".markdown"):
        content = _load_txt(path)  # Markdown 作为纯文本处理，切块时保留原始格式
    elif ext == ".pdf":
        content = _load_pdf(path)
    elif ext == ".docx":
        content = _load_docx(path)
    elif ext == ".csv":
        content = _load_csv(path)
    elif ext == ".json":
        content = _load_json(path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")

    if not content.strip():
        raise ValueError(f"文件内容为空: {file_path}")

    return Document(
        id=doc_id,
        content=content,
        source=str(path),
        filename=path.name,
        file_type=file_type,
        size=size,
        metadata={"ext": ext},
    )


def _load_txt(path: Path) -> str:
    """加载纯文本文件（自动尝试 UTF-8 / GBK）。"""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法解码文件: {path}")


def _load_pdf(path: Path) -> str:
    """加载 PDF 文件（使用 PyPDF2）。"""
    try:
        import PyPDF2
    except ImportError:
        raise ValueError("PyPDF2 未安装，无法解析 PDF。请安装: pip install PyPDF2")

    text_parts = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _load_docx(path: Path) -> str:
    """加载 DOCX 文件（使用 python-docx）。"""
    try:
        import docx
    except ImportError:
        raise ValueError("python-docx 未安装，无法解析 DOCX。请安装: pip install python-docx")

    doc = docx.Document(str(path))
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(text_parts)


def _load_csv(path: Path) -> str:
    """加载 CSV 文件，转换为可读文本格式。"""
    rows = []
    with open(path, "r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)


def _load_json(path: Path) -> str:
    """加载 JSON 文件，转换为可读文本格式。"""
    data = json.loads(path.read_text(encoding="utf-8"))
    return json.dumps(data, ensure_ascii=False, indent=2)


def load_text_from_url(url: str, markdown: str) -> Document:
    """从 URL 提取的文本创建 Document 对象。

    Args:
        url: 原始 URL
        markdown: 提取的 Markdown 文本

    Returns:
        Document 对象
    """
    from urllib.parse import urlparse

    parsed = urlparse(url)
    filename = parsed.path.rsplit("/", 1)[-1] or "index"
    doc_id = f"url_{filename}"

    return Document(
        id=doc_id,
        content=markdown,
        source=url,
        filename=filename,
        file_type="url",
        size=len(markdown.encode("utf-8")),
        metadata={"url": url, "imported_via": "url"},
    )
